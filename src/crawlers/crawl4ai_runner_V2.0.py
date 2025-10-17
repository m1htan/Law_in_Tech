# src/crawlers/crawl4ai_runner.py
import re
import time
import csv
import json
import hashlib
import logging
from logging.handlers import RotatingFileHandler
from pathlib import Path
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional, Set
from urllib.parse import urljoin, urlparse
import asyncio
import urllib.robotparser as robotparser
import requests
import pdfplumber
from bs4 import BeautifulSoup
import shutil
import subprocess

from playwright.async_api import async_playwright, BrowserContext, Page

# =============================
# CONFIG
# =============================
BASE = Path(__file__).resolve().parents[2]
OUT_PDF_DIR = BASE / "outputs" / "raw" / "pdf"
OUT_TXT_DIR = BASE / "outputs" / "raw" / "txt"
OUT_LOGS   = BASE / "outputs" / "logs"
OUT_CSV    = BASE / "outputs" / "raw" / "csv" / "all.csv"

for d in [OUT_PDF_DIR, OUT_TXT_DIR, OUT_LOGS, OUT_CSV.parent]:
    d.mkdir(parents=True, exist_ok=True)

DATE_MIN = datetime(2022, 1, 1, tzinfo=timezone.utc)
LOG_FILE = OUT_LOGS / "crawl.log"

# Crawl tuning
REQUEST_DELAY = 1.2   # delay giữa các trang/nhấp
WAIT_NET      = "networkidle"  # điều kiện chờ Playwright
WAIT_MS       = 1200           # chờ thêm sau AJAX
MAX_QH_PAGES  = 200            # upper bound an toàn cho QH mỗi tab/type
MAX_CP_PAGES  = 500            # upper bound an toàn cho Chinhphu

UA = ("Mozilla/5.0 (Macintosh; Intel Mac OS X 15_6_1) "
      "AppleWebKit/537.36 (KHTML, like Gecko) "
      "Chrome/141.0.0.0 Safari/537.36")

# Timeout/Retry
NAV_TIMEOUT_MS   = 90_000        # 90s cho điều hướng nặng (QH detail)
SEL_TIMEOUT_MS   = 45_000        # 45s chờ selector
RETRY_ATTEMPTS   = 3
RETRY_BACKOFF_S  = 2.0

# domains ồn, chặn để vào trang nhanh hơn
BLOCKED_RES = [
    "googletagmanager.com", "google-analytics.com", "g.doubleclick.net",
    "analytics.google.com", "fonts.googleapis.com", "fonts.gstatic.com",
]

# HTTP session với retry
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

def make_http_session() -> requests.Session:
    s = requests.Session()
    retries = Retry(
        total=4, backoff_factor=0.8,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods=["GET", "HEAD"]
    )
    s.mount("http://", HTTPAdapter(max_retries=retries))
    s.mount("https://", HTTPAdapter(max_retries=retries))
    s.headers.update({"User-Agent": UA})
    return s

HTTP = make_http_session()

# =============================
# LOGGER
# =============================
logger = logging.getLogger("crawler")
logger.setLevel(logging.INFO)
fmt = logging.Formatter("%(asctime)s [%(levelname)s] %(name)s - %(message)s")
fh = RotatingFileHandler(LOG_FILE, maxBytes=10_000_000, backupCount=3, encoding="utf-8")
fh.setFormatter(fmt)
sh = logging.StreamHandler()
sh.setFormatter(fmt)
logger.handlers.clear()
logger.addHandler(fh)
logger.addHandler(sh)

# =============================
# HELPERS
# =============================
def sha1_bytes(b: bytes) -> str:
    h = hashlib.sha1(); h.update(b); return h.hexdigest()

def pdf_to_text(pdf_path: Path) -> str:
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
    except Exception as e:
        logger.warning(f"PDF to text failed for {pdf_path.name}: {e}")
        return ""

def docx_to_text(docx_path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(docx_path))
        parts = []
        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)
        # bảng (table) – lấy text từng ô (tuỳ file có thể bỏ)
        for table in doc.tables:
            for row in table.rows:
                row_txt = "\t".join(cell.text.strip() for cell in row.cells)
                if row_txt.strip():
                    parts.append(row_txt)
        text = "\n".join(parts)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()
    except Exception as e:
        logger.warning(f"DOCX to text failed for {docx_path.name}: {e}")
        return ""

def doc_to_text_via_antiword(doc_path: Path) -> Optional[str]:
    """Dùng antiword nếu có."""
    if shutil.which("antiword"):
        try:
            res = subprocess.run(
                ["antiword", "-m", "UTF-8.txt", str(doc_path)],
                stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True
            )
            return res.stdout.decode("utf-8", errors="ignore").strip()
        except Exception as e:
            logger.warning(f"antiword failed for {doc_path.name}: {e}")
    return None

def doc_to_text_via_catdoc(doc_path: Path) -> Optional[str]:
    """Dùng catdoc nếu có."""
    if shutil.which("catdoc"):
        try:
            res = subprocess.run(
                ["catdoc", "-w", str(doc_path)],
                stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True
            )
            return res.stdout.decode("utf-8", errors="ignore").strip()
        except Exception as e:
            logger.warning(f"catdoc failed for {doc_path.name}: {e}")
    return None

def doc_to_text_via_libreoffice(doc_path: Path) -> Optional[str]:
    """
    Dự phòng cuối: dùng LibreOffice để convert .doc -> .txt.
    Tạo file .txt cạnh file gốc rồi đọc vào.
    """
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    try:
        out_dir = doc_path.parent
        # convert ra .txt (Text) – trên macOS/LO thường dùng filter "Text"
        subprocess.run(
            [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", str(out_dir), str(doc_path)],
            stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True
        )
        txt_path = doc_path.with_suffix(".txt")
        if txt_path.exists():
            return txt_path.read_text(encoding="utf-8", errors="ignore").strip()
    except Exception as e:
        logger.warning(f"LibreOffice convert-to txt failed for {doc_path.name}: {e}")
    return None

def doc_to_text(doc_path: Path) -> str:
    """
    Chiến lược: antiword -> catdoc -> (dự phòng) LibreOffice.
    """
    for fn in (doc_to_text_via_antiword, doc_to_text_via_catdoc, doc_to_text_via_libreoffice):
        txt = fn(doc_path)
        if txt:
            return txt
    logger.warning(f"No available tool to extract .doc for {doc_path.name}")
    return ""

def extract_text_generic(local_path: Path) -> str:
    suf = local_path.suffix.lower()
    if suf == ".pdf":
        return pdf_to_text(local_path)
    if suf == ".docx":
        return docx_to_text(local_path)
    if suf == ".doc":
        return doc_to_text(local_path)
    # có thể mở rộng .rtf/.zip... nếu cần
    return ""

def append_csv(rec: Dict[str, Any], path: Path):
    keys = [
        "source_label", "list_url", "detail_url", "download_url",
        "pdf_local", "txt_local", "sha1_pdf", "pdf_text_len",
        "title", "crawl_time"
    ]
    new_file = not path.exists()
    with path.open("a", encoding="utf-8", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=keys)
        if new_file:
            writer.writeheader()
        writer.writerow({k: rec.get(k, "") for k in keys})

def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()

def abs_url(base: str, href: str) -> str:
    return urljoin(base, href)

def dedup_order(seq: List[str]) -> List[str]:
    seen = set(); out = []
    for x in seq:
        if x not in seen:
            seen.add(x); out.append(x)
    return out

def extract_pdf_links(html: str, base_url: str) -> List[str]:
    soup = BeautifulSoup(html, "lxml")
    pdfs = []
    for a in soup.select("a[href]"):
        href = a["href"].strip()
        if re.search(r"\.pdf(\?|$)", href, re.I):
            pdfs.append(abs_url(base_url, href))
    return dedup_order(pdfs)

def is_pdf_bytes(b: bytes) -> bool:
    # PDF luôn bắt đầu bằng "%PDF-"
    return b.startswith(b"%PDF-")

# ====== ROBOTS ======
ROBOTS_CACHE: Dict[str, robotparser.RobotFileParser] = {}

def robots_allow(url: str) -> bool:
    """Enforce robots.txt (đang chỉ dùng cho QH vì bạn yêu cầu)."""
    parsed = urlparse(url)
    base = f"{parsed.scheme}://{parsed.netloc}"
    rp = ROBOTS_CACHE.get(base)
    if rp is None:
        rp = robotparser.RobotFileParser()
        rp.set_url(urljoin(base, "/robots.txt"))
        try:
            rp.read()
        except Exception as e:
            logger.warning(f"[ROBOTS] fetch failed {base}: {e} -> disallow")
            ROBOTS_CACHE[base] = rp
            return False
        ROBOTS_CACHE[base] = rp
    try:
        return rp.can_fetch(UA, url)
    except Exception:
        return False

# ====== Missing helpers used by main() ======
def save_pdf(bytes_data: bytes) -> tuple[str, str]:
    """Lưu bytes thành .pdf (giữ nguyên hành vi main() hiện tại)."""
    sha1 = sha1_bytes(bytes_data)
    p = OUT_PDF_DIR / f"{sha1}.pdf"
    if not p.exists():
        p.write_bytes(bytes_data)
        logger.info(f"Saved PDF: {p}")
    else:
        logger.info(f"PDF exists (dedup): {p}")
    return str(p), sha1

def not_seen(sha1: str, seen: Set[str]) -> bool:
    if sha1 in seen:
        return False
    seen.add(sha1)
    return True

def http_download_to_path(url: str, out_dir: Path, referer: Optional[str] = None, insecure: bool = False) -> Optional[Path]:
    """Tải bằng requests rồi lưu file, trả về Path. insecure=True => verify=False (dành cho QH gateway)."""
    out_dir.mkdir(parents=True, exist_ok=True)
    try:
        headers = {"User-Agent": UA}
        if referer:
            headers["Referer"] = referer
        r = HTTP.get(url, headers=headers, timeout=60, verify=not insecure)
        if r.status_code != 200 or not r.content:
            logger.warning(f"HTTP {r.status_code} on {url}")
            return None
        name = Path(urlparse(url).path).name or "download.bin"
        dest = out_dir / name
        dest.write_bytes(r.content)
        return dest
    except Exception as e:
        logger.warning(f"requests download failed {url}: {e}")
        return None

def is_file_url(u: str) -> bool:
    return bool(re.search(r"\.(pdf|docx?|zip)$", u, re.I))

# ============ HTTP download (không dùng Crawl4AI .download) ============
def http_download(url: str, referer: Optional[str] = None, timeout: int = 60) -> Optional[bytes]:
    try:
        headers = {}
        if referer:
            headers["Referer"] = referer
        r = HTTP.get(url, headers=headers, timeout=timeout)
        if r.status_code == 200 and r.content:
            return r.content
        logger.warning(f"HTTP {r.status_code} on {url}")
        return None
    except Exception as e:
        logger.warning(f"requests download failed {url}: {e}")
        return None

# =============================
# QH — gọi trực tiếp loadPagingAjax(...) trong trang
# =============================
def _qh_type_from_url(u: str) -> int:
    if "du-thao-luat" in u: return 1
    if "du-thao-nghi-quyet" in u: return 2
    return 3  # pháp lệnh

async def qh_list_detail_urls(ctx: BrowserContext, list_url: str) -> List[str]:
    base = "{uri.scheme}://{uri.netloc}".format(uri=urlparse(list_url))
    t = _qh_type_from_url(list_url)
    details: List[str] = []
    page: Page = await ctx.new_page()
    try:
        await page.goto(list_url, wait_until=WAIT_NET)
        await page.wait_for_timeout(WAIT_MS)
        # Hai tab: nav-profile (TrangThai=0), nav-contact (TrangThai=1)
        tabs = [("nav-profile", 0), ("nav-contact", 1)]
        for container, trang_thai in tabs:
            empty_hits = 0
            for p in range(0, MAX_QH_PAGES):
                # truyền tham số qua một object duy nhất (đúng chuẩn evaluate)
                js = """
                  (args) => {
                    const {p, containerId, t, trangThai} = args;
                    const q = `?handler=DanhSachDuThao&PageSize=10&ContainerBindData=${containerId}&Type=${t}&TrangThai=${trangThai}&pageNumber=${p}`;
                    if (window.loadPagingAjax) {
                      window.loadPagingAjax(p, 10, 'DanhSachDuThao', containerId, q);
                    } else {
                      const u = new URL(window.location.href);
                      window.location.href = u.pathname + q;
                    }
                  }
                """
                await page.evaluate(js, {"p": p, "containerId": container, "t": t, "trangThai": trang_thai})
                await page.wait_for_timeout(WAIT_MS)

                # lấy HTML của container hiện tại
                try:
                    html = await page.inner_html(f"#{container}")
                except Exception:
                    html = await page.content()

                soup = BeautifulSoup(html, "lxml")
                cards = []
                for a in soup.select("a.d-inline-block[href]"):
                    href = a["href"].strip()
                    if href.startswith("/dt/"):
                        cards.append(abs_url(base, href))
                cards = dedup_order(cards)
                new_cards = [c for c in cards if c not in details]
                logger.info(f"[QH] container={container} tt={trang_thai} page={p} cards={len(cards)} new={len(new_cards)}")

                if not new_cards:
                    empty_hits += 1
                else:
                    details.extend(new_cards)
                    empty_hits = 0

                if empty_hits >= 2:
                    break
                await page.wait_for_timeout(int(REQUEST_DELAY * 1000))
    finally:
        await page.close()
    return dedup_order(details)

async def qh_detail_files(ctx: BrowserContext, detail_url: str) -> List[str]:
    if is_file_url(detail_url):
        return [detail_url]

    sel_candidates = [
        "a:has-text('Tải file đính kèm')",
        "a i.fa-download",                 # icon tải đính kèm
        "a[href*='uploadFiles']",
    ]

    for attempt in range(1, RETRY_ATTEMPTS + 1):
        page = await ctx.new_page()
        try:
            await page.goto(detail_url, wait_until="domcontentloaded")
            # đợi container tab chính hiện diện hoặc phần đính kèm
            try:
                await page.wait_for_selector("#nav-tabContent, .tab-content", timeout=SEL_TIMEOUT_MS)
            except Exception:
                pass

            # cố gắng bấm đúng tab có đính kèm (nếu có)
            try:
                # nhiều trang có anchor nav-ykien/nav-home/nav-profile, ta không cần bấm nếu đã có link tải
                pass
            except Exception:
                pass

            # đợi thêm chút để AJAX vẽ danh sách file
            await page.wait_for_timeout(WAIT_MS)

            html = await page.content()
            soup = BeautifulSoup(html, "lxml")
            out = []
            for a in soup.select("a[href]"):
                href = a.get("href", "").strip()
                if re.search(r"\.(pdf|docx?)($|\?)", href, re.I) or "uploadFiles" in href:
                    out.append(abs_url(detail_url, href))
            out = dedup_order(out)
            if out:
                return out

            # nếu chưa thấy, thử click dùng selector “Tải file đính kèm”
            clicked = False
            for sel in sel_candidates:
                try:
                    el = await page.query_selector(sel)
                    if el:
                        await el.click()
                        clicked = True
                        break
                except Exception:
                    pass
            if clicked:
                await page.wait_for_timeout(WAIT_MS)
                html = await page.content()
                soup = BeautifulSoup(html, "lxml")
                out = []
                for a in soup.select("a[href]"):
                    href = a.get("href", "").strip()
                    if re.search(r"\.(pdf|docx?)($|\?)", href, re.I) or "uploadFiles" in href:
                        out.append(abs_url(detail_url, href))
                out = dedup_order(out)
                if out:
                    return out

            # không thấy file → thử lần sau
            logger.warning(f"[QH] no files found on {detail_url} (attempt {attempt})")
        except Exception as e:
            logger.warning(f"[QH] detail fetch failed {detail_url} (attempt {attempt}): {e}")
        finally:
            await page.close()
        await asyncio.sleep(RETRY_BACKOFF_S * attempt)

    return []

# =============================
# CHÍNH PHỦ — duyệt ?page=N và/hoặc click phân trang
# =============================
async def cp_list_detail_urls(ctx: BrowserContext, list_url: str) -> List[str]:
    details: List[str] = []
    page = await ctx.new_page()
    try:
        empty_hits = 0
        for p in range(1, MAX_CP_PAGES + 1):
            url = f"{list_url}?page={p}"
            await page.goto(url, wait_until=WAIT_NET)
            await page.wait_for_timeout(WAIT_MS)
            html = await page.content()
            soup = BeautifulSoup(html, "lxml")
            links = []
            for a in soup.select("a[href]"):
                href = a["href"].strip()
                full = abs_url(list_url, href)
                # chỉ lấy TRANG CHI TIẾT trên chinhphu.vn, bỏ link file ở datafiles.*
                if "chinhphu.vn" in urlparse(full).netloc and "/du-thao" in full and not is_file_url(full):
                    links.append(full)
            links = dedup_order(links)
            new_links = [u for u in links if u not in details]
            logger.info(f"[CP] page={p} cards={len(links)} new={len(new_links)}")
            if not new_links:
                empty_hits += 1
            else:
                details.extend(new_links)
                empty_hits = 0
            if empty_hits >= 2:
                break
            await page.wait_for_timeout(int(REQUEST_DELAY * 1000))
    finally:
        await page.close()
    return dedup_order(details)

async def cp_detail_files(ctx: BrowserContext, detail_url: str) -> List[str]:
    if is_file_url(detail_url):
        return [detail_url]
    for attempt in range(1, RETRY_ATTEMPTS + 1):
        page = await ctx.new_page()
        try:
            await page.goto(detail_url, wait_until="domcontentloaded")
            await page.wait_for_timeout(WAIT_MS)
            html = await page.content()
            soup = BeautifulSoup(html, "lxml")
            out = []
            for a in soup.select("a[href]"):
                href = a.get("href", "").strip()
                if re.search(r"\.(pdf|docx?)($|\?)", href, re.I):
                    out.append(abs_url(detail_url, href))
            out = dedup_order(out)
            if out:
                return out
        except Exception as e:
            logger.warning(f"[CP] detail fetch failed {detail_url} (attempt {attempt}): {e}")
        finally:
            await page.close()
        await asyncio.sleep(RETRY_BACKOFF_S * attempt)
    return []

# =============================
# MST — lấy “Xem chi tiết” rồi tab “Văn bản gốc/PDF”
# =============================
async def mst_list_detail_urls(ctx: BrowserContext, list_url: str) -> List[str]:
    details: List[str] = []
    page = await ctx.new_page()
    try:
        await page.goto(list_url, wait_until=WAIT_NET)
        await page.wait_for_timeout(WAIT_MS)
        html = await page.content()
        soup = BeautifulSoup(html, "lxml")
        for a in soup.select("a.view-more[href]"):
            href = a["href"].strip()
            if href and not href.lower().startswith("javascript"):
                details.append(urljoin(list_url, href))
    finally:
        await page.close()
    details = dedup_order(details)
    logger.info(f"[MST] collected detail urls: {len(details)}")
    return details

async def mst_detail_files(ctx: BrowserContext, detail_url: str) -> List[str]:
    if is_file_url(detail_url):
        return [detail_url]
    for attempt in range(1, RETRY_ATTEMPTS + 1):
        page = await ctx.new_page()
        try:
            await page.goto(detail_url, wait_until="domcontentloaded")
            await page.wait_for_timeout(WAIT_MS)
            # click tab Văn bản gốc/PDF nếu có thật (href không phải javascript)
            try:
                candidates = await page.locator("a", has_text=re.compile("Văn bản gốc/PDF", re.I)).all()
                for el in candidates:
                    href = await el.get_attribute("href")
                    if href and not href.lower().startswith("javascript"):
                        await el.click()
                        await page.wait_for_timeout(WAIT_MS)
                        break
            except Exception:
                pass

            html = await page.content()
            soup = BeautifulSoup(html, "lxml")
            links = []
            for a in soup.select("a.doc-download[href], a[href$='.pdf'], a[href$='.doc'], a[href$='.docx']"):
                href = a.get("href", "").strip()
                if href and not href.lower().startswith("javascript"):
                    links.append(urljoin(detail_url, href))
            links = dedup_order(links)
            if links:
                return links
        except Exception as e:
            logger.warning(f"[MST] detail fetch failed {detail_url} (attempt {attempt}): {e}")
        finally:
            await page.close()
        await asyncio.sleep(RETRY_BACKOFF_S * attempt)
    return []

# =============================
# CORE: DOWNLOAD & SAVE
# =============================
def save_binary_as_pdfname(bytes_data: bytes, prefer_pdf: bool, src_url: str) -> tuple[str, str]:
    file_sha1 = sha1_bytes(bytes_data)
    ext = ".pdf" if prefer_pdf else (".pdf" if src_url.lower().endswith(".pdf") else Path(urlparse(src_url).path).suffix or ".bin")
    p = OUT_PDF_DIR / f"{file_sha1}{ext}"
    if not p.exists():
        p.write_bytes(bytes_data)
        logger.info(f"Saved BIN: {p}")
    else:
        logger.info(f"BIN exists (dedup): {p}")
    return str(p), file_sha1

def handle_download_and_save(source_label, list_url, detail_url, download_url, seen_sha1: Set[str]):
    bin_data = http_download(download_url, referer=detail_url)
    if not bin_data:
        logger.warning(f"Failed to download: {download_url}")
        return
    prefer_pdf = download_url.lower().endswith(".pdf")
    file_path, sha1 = save_binary_as_pdfname(bin_data, prefer_pdf, download_url)
    if sha1 in seen_sha1:
        return
    seen_sha1.add(sha1)

    txt = ""
    if prefer_pdf:
        txt = pdf_to_text(Path(file_path))
    txt_path = OUT_TXT_DIR / f"{sha1}.txt"
    try:
        txt_path.write_text(txt or "", encoding="utf-8")
    except Exception as e:
        logger.warning(f"Write txt failed {txt_path}: {e}")

    rec = {
        "source_label": source_label,
        "list_url": list_url,
        "detail_url": detail_url,
        "download_url": download_url,
        "pdf_local": str(file_path),      # có thể là .pdf/.doc/.docx
        "txt_local": str(txt_path),
        "sha1_pdf": sha1,
        "pdf_text_len": len(txt or ""),
        "title": "",
        "crawl_time": now_iso(),
    }
    append_csv(rec, OUT_CSV)

from urllib.parse import urlparse

ALLOWED_INSECURE_HOSTS = {"gatewayduthaoonline.quochoi.vn"}  # chỉ QH gateway

async def playwright_download(context, url: str, out_dir: Path, timeout_ms: int = 60000) -> Optional[Path]:
    """Tải file bằng Playwright để tránh lỗi SSL của requests."""
    out_dir.mkdir(parents=True, exist_ok=True)
    page = await context.new_page()
    try:
        # Quan trọng: expect_download bao quanh hành động gây ra download
        async with page.expect_download(timeout=timeout_ms):
            await page.goto(url, wait_until="domcontentloaded")
        download = await page.wait_for_event("download", timeout=timeout_ms)
        suggested = download.suggested_filename or "download.bin"
        dest = out_dir / suggested
        await download.save_as(str(dest))
        return dest
    finally:
        await page.close()

# =============================
# MAIN RUNNER
# =============================
async def main():
    logger.info("=== START RUN ===")
    seen = set()
    browser = await async_playwright().start()
    context = await browser.chromium.launch_persistent_context(
        user_data_dir="/tmp/qh_ctx",
        headless=True,
        ignore_https_errors=True,  # cho phép QH SSL lỗi nhẹ
        args=["--disable-gpu", "--no-sandbox"],
    )

    # --- DOMAIN CONFIG ---
    qh_luat = "https://duthaoonline.quochoi.vn/du-thao/du-thao-luat"
    qh_nq = "https://duthaoonline.quochoi.vn/du-thao/du-thao-nghi-quyet"
    qh_pl = "https://duthaoonline.quochoi.vn/du-thao/du-thao-phap-lenh"
    cp_url = "https://chinhphu.vn/du-thao-vbqppl"
    mst_url = "https://mst.gov.vn/van-ban-phap-luat.htm"

    # --- QH: LUẬT ---
    try:
        logger.info(f":: Crawl QH list begin [DU_THAO_QH_LUAT] {qh_luat}")
        if robots_allow(qh_luat):
            detail_urls = await qh_list_detail_urls(context, qh_luat)
            logger.info(f"[DU_THAO_QH_LUAT] detail urls = {len(detail_urls)}")

            for du in detail_urls:
                if not robots_allow(du):
                    logger.warning(f"[ROBOTS] skip detail: {du}")
                    continue

                files = await qh_detail_files(context, du)
                for download_url in files:
                    host = urlparse(download_url).netloc

                    saved_path = http_download_to_path(
                        download_url,
                        OUT_PDF_DIR,
                        referer=du,
                        insecure=(host in ALLOWED_INSECURE_HOSTS),
                    )
                    if not saved_path:
                        logger.warning(f"Failed to download: {download_url}")
                        continue

                    bin_bytes = Path(saved_path).read_bytes()
                    is_pdf = is_pdf_bytes(bin_bytes)
                    prefer_pdf = is_pdf or download_url.lower().endswith(".pdf")

                    file_path, sha1 = save_binary_as_pdfname(bin_bytes, prefer_pdf, download_url)
                    if not not_seen(sha1, seen):
                        continue

                    txt = extract_text_generic(Path(file_path))
                    txt_path = OUT_TXT_DIR / f"{sha1}.txt"
                    txt_path.write_text(txt or "", encoding="utf-8")

                    rec = {
                        "source_label": "DU_THAO_QH_LUAT",
                        "list_url": qh_luat,
                        "detail_url": du,
                        "download_url": download_url,
                        "pdf_local": str(file_path),
                        "txt_local": str(txt_path),
                        "sha1_pdf": sha1,
                        "pdf_text_len": len(txt or ""),
                        "title": Path(saved_path).name,
                        "crawl_time": now_iso(),
                    }
                    append_csv(rec, OUT_CSV)
        else:
            logger.warning(f"[ROBOTS] Disallowed: {qh_luat}")
    except Exception as e:
        logger.exception(f"DU_THAO_QH_LUAT failed: {e}")
    logger.info(":: Crawl QH list end [DU_THAO_QH_LUAT]")

    # --- QH: NGHỊ QUYẾT ---
    try:
        logger.info(f":: Crawl QH list begin [DU_THAO_QH_NGHI_QUYET] {qh_nq}")
        if robots_allow(qh_nq):
            detail_urls = await qh_list_detail_urls(context, qh_nq)
            logger.info(f"[DU_THAO_QH_NGHI_QUYET] detail urls = {len(detail_urls)}")

            for du in detail_urls:
                if not robots_allow(du):
                    continue

                files = await qh_detail_files(context, du)
                for download_url in files:
                    host = urlparse(download_url).netloc

                    saved_path = http_download_to_path(
                        download_url,
                        OUT_PDF_DIR,
                        referer=du,
                        insecure=(host in ALLOWED_INSECURE_HOSTS),
                    )
                    if not saved_path:
                        logger.warning(f"Failed to download: {download_url}")
                        continue

                    bin_bytes = Path(saved_path).read_bytes()
                    is_pdf = is_pdf_bytes(bin_bytes)
                    prefer_pdf = is_pdf or download_url.lower().endswith(".pdf")

                    file_path, sha1 = save_binary_as_pdfname(bin_bytes, prefer_pdf, download_url)
                    if not not_seen(sha1, seen):
                        continue

                    txt = extract_text_generic(Path(file_path))
                    txt_path = OUT_TXT_DIR / f"{sha1}.txt"
                    txt_path.write_text(txt or "", encoding="utf-8")

                    rec = {
                        "source_label": "DU_THAO_QH_NGHI_QUYET",
                        "list_url": qh_nq,
                        "detail_url": du,
                        "download_url": download_url,
                        "pdf_local": str(file_path),
                        "txt_local": str(txt_path),
                        "sha1_pdf": sha1,
                        "pdf_text_len": len(txt or ""),
                        "title": Path(saved_path).name,
                        "crawl_time": now_iso(),
                    }
                    append_csv(rec, OUT_CSV)
        else:
            logger.warning(f"[ROBOTS] Disallowed: {qh_nq}")
    except Exception as e:
        logger.exception(f"DU_THAO_QH_NGHI_QUYET failed: {e}")
    logger.info(":: Crawl QH list end [DU_THAO_QH_NGHI_QUYET]")

    # --- QH: PHÁP LỆNH ---
    try:
        logger.info(f":: Crawl QH list begin [DU_THAO_QH_PHAP_LENH] {qh_pl}")
        if robots_allow(qh_pl):
            detail_urls = await qh_list_detail_urls(context, qh_pl)
            logger.info(f"[DU_THAO_QH_PHAP_LENH] detail urls = {len(detail_urls)}")

            for du in detail_urls:
                if not robots_allow(du):
                    continue

                files = await qh_detail_files(context, du)
                for download_url in files:
                    host = urlparse(download_url).netloc

                    saved_path = http_download_to_path(
                        download_url,
                        OUT_PDF_DIR,
                        referer=du,
                        insecure=(host in ALLOWED_INSECURE_HOSTS),
                    )
                    if not saved_path:
                        logger.warning(f"Failed to download: {download_url}")
                        continue

                    bin_bytes = Path(saved_path).read_bytes()
                    is_pdf = is_pdf_bytes(bin_bytes)
                    prefer_pdf = is_pdf or download_url.lower().endswith(".pdf")

                    file_path, sha1 = save_binary_as_pdfname(bin_bytes, prefer_pdf, download_url)
                    if not not_seen(sha1, seen):
                        continue

                    txt = extract_text_generic(Path(file_path))
                    txt_path = OUT_TXT_DIR / f"{sha1}.txt"
                    txt_path.write_text(txt or "", encoding="utf-8")

                    rec = {
                        "source_label": "DU_THAO_QH_PHAP_LENH",
                        "list_url": qh_pl,
                        "detail_url": du,
                        "download_url": download_url,
                        "pdf_local": str(file_path),
                        "txt_local": str(txt_path),
                        "sha1_pdf": sha1,
                        "pdf_text_len": len(txt or ""),
                        "title": Path(saved_path).name,
                        "crawl_time": now_iso(),
                    }
                    append_csv(rec, OUT_CSV)
        else:
            logger.warning(f"[ROBOTS] Disallowed: {qh_pl}")
    except Exception as e:
        logger.exception(f"DU_THAO_QH_PHAP_LENH failed: {e}")
    logger.info(":: Crawl QH list end [DU_THAO_QH_PHAP_LENH]")

    # --- CHÍNH PHỦ ---
    try:
        logger.info(f":: Crawl CP list begin [DU_THAO_CP] {cp_url}")
        detail_urls = await cp_list_detail_urls(context, cp_url)
        logger.info(f"[DU_THAO_CP] detail urls = {len(detail_urls)}")
        for du in detail_urls:
            files = await cp_detail_files(context, du)
            for download_url in files:
                host = urlparse(download_url).netloc

                saved_path = http_download_to_path(
                    download_url,
                    OUT_PDF_DIR,
                    referer=du,
                    insecure=(host in ALLOWED_INSECURE_HOSTS),
                )
                if not saved_path:
                    logger.warning(f"Failed to download: {download_url}")
                    continue

                bin_bytes = Path(saved_path).read_bytes()
                is_pdf = is_pdf_bytes(bin_bytes)
                prefer_pdf = is_pdf or download_url.lower().endswith(".pdf")

                file_path, sha1 = save_binary_as_pdfname(bin_bytes, prefer_pdf, download_url)
                if not not_seen(sha1, seen):
                    continue

                txt = extract_text_generic(Path(file_path))
                txt_path = OUT_TXT_DIR / f"{sha1}.txt"
                txt_path.write_text(txt or "", encoding="utf-8")

                rec = {
                    "source_label": "DU_THAO_CP",
                    "list_url": cp_url,
                    "detail_url": du,
                    "download_url": download_url,
                    "pdf_local": str(file_path),
                    "txt_local": str(txt_path),
                    "sha1_pdf": sha1,
                    "pdf_text_len": len(txt or ""),
                    "title": Path(saved_path).name,
                    "crawl_time": now_iso(),
                }
                append_csv(rec, OUT_CSV)
    except Exception as e:
        logger.exception(f"DU_THAO_CP failed: {e}")
    logger.info(":: Crawl CP list end [DU_THAO_CP]")

    # --- MST ---
    try:
        logger.info(f":: Crawl MST list begin [MST] {mst_url}")
        detail_urls = await mst_list_detail_urls(context, mst_url)
        logger.info(f"[MST] detail urls = {len(detail_urls)}")
        for du in detail_urls:
            files = await mst_detail_files(context, du)
            for download_url in files:
                host = urlparse(download_url).netloc

                saved_path = http_download_to_path(
                    download_url,
                    OUT_PDF_DIR,
                    referer=du,
                    insecure=(host in ALLOWED_INSECURE_HOSTS),
                )
                if not saved_path:
                    logger.warning(f"Failed to download: {download_url}")
                    continue

                bin_bytes = Path(saved_path).read_bytes()
                is_pdf = is_pdf_bytes(bin_bytes)
                prefer_pdf = is_pdf or download_url.lower().endswith(".pdf")

                file_path, sha1 = save_binary_as_pdfname(bin_bytes, prefer_pdf, download_url)
                if not not_seen(sha1, seen):
                    continue

                txt = extract_text_generic(Path(file_path))
                txt_path = OUT_TXT_DIR / f"{sha1}.txt"
                txt_path.write_text(txt or "", encoding="utf-8")

                rec = {
                    "source_label": "MST",
                    "list_url": mst_url,
                    "detail_url": du,
                    "download_url": download_url,
                    "pdf_local": str(file_path),
                    "txt_local": str(txt_path),
                    "sha1_pdf": sha1,
                    "pdf_text_len": len(txt or ""),
                    "title": Path(saved_path).name,
                    "crawl_time": now_iso(),
                }
                append_csv(rec, OUT_CSV)
    except Exception as e:
        logger.exception(f"MST failed: {e}")
    logger.info(":: Crawl MST list end [MST]")

    await context.close()
    await browser.stop()
    logger.info("=== END RUN ===")

if __name__ == "__main__":
    import asyncio
    logger.info("=== START RUN ===")
    try:
        asyncio.run(main())
    finally:
        logger.info("=== END RUN ===")
